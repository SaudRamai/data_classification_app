"""
Policy & Guidance page for the Data Governance Application.
- Decision Tree and Quick Reference checklist derived from policy document.
- Links classification actions to rationale capture and exceptions.
"""
from src.ui.quick_links import render_quick_links
import sys
import os
import pathlib

# Add the project root to the Python path
_here = pathlib.Path(str(__file__)).resolve()
_dir = _here.parent
# Traverse up to find directory containing 'src'
for _ in range(3):
    if (_dir / "src").exists():
        if str(_dir) not in sys.path:
            sys.path.insert(0, str(_dir))
        break
    _dir = _dir.parent

import streamlit as st

# MUST be the first Streamlit command
st.set_page_config(page_title="Handling Rules & Policy Guidance", page_icon="📘", layout="wide")

import pandas as pd
import time
import io
try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx
except ImportError:
    docx = None

import base64
from src.ui.theme import apply_global_theme
from src.services.compliance_service import compliance_service as _ncs
from src.connectors.snowflake_connector import snowflake_connector
from src.config.settings import settings
from src.components.filters import render_global_filters

from src.services.authorization_service import authz

# Apply centralized theme
apply_global_theme()

# ============================================================================
# RBAC CHECK
# ============================================================================
try:
    _ident = authz.get_current_identity()
    if not authz.is_consumer(_ident):
        if authz._is_bypass():
            st.warning("RBAC bypass active (testing). UI is visible for verification.")
        else:
            st.error("You do not have permission to access the Policy module.")
            st.stop()
except Exception as _auth_err:
    if not authz._is_bypass():
        st.warning(f"Authorization check failed: {_auth_err}")
        st.stop()

# Global Filters

with st.sidebar:
    g_filters = render_global_filters(key_prefix="policy")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&family=Outfit:wght@500;600;700&display=swap');

.stApp {
    font-family: 'Inter', sans-serif;
}

/* Premium Glassmorphism Token */
:root {
    --glass-bg: rgba(30, 41, 59, 0.4);
    --glass-border: rgba(255, 255, 255, 0.08);
    --accent-blue: #3b82f6;
    --accent-green: #10b981;
}

/* Page Hero Upgrade */
.page-hero {
    background: linear-gradient(135deg, rgba(30, 41, 59, 0.8) 0%, rgba(15, 23, 42, 0.9) 100%);
    padding: 2.5rem;
    border-radius: 24px;
    border: 1px solid var(--glass-border);
    margin-bottom: 2.5rem;
    position: relative;
    overflow: hidden;
}

.page-hero::before {
    content: '';
    position: absolute;
    top: -50%;
    right: -10%;
    width: 300px;
    height: 300px;
    background: radial-gradient(circle, rgba(59, 130, 246, 0.1) 0%, transparent 70%);
    z-index: 0;
}

.hero-title {
    font-family: 'Outfit', sans-serif;
    font-size: 2.5rem;
    font-weight: 700;
    margin: 0;
    background: linear-gradient(to right, #fff, #94a3b8);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
}

.hero-subtitle {
    color: #94a3b8;
    font-size: 1.1rem;
    max-width: 800px;
    margin-top: 0.5rem;
}

/* Section Headers */
.section-header {
    font-family: 'Outfit', sans-serif;
    font-size: 0.85rem;
    letter-spacing: 0.15em;
    text-transform: uppercase;
    color: #64748b;
    font-weight: 700;
    margin: 2.5rem 0 1.5rem 0;
    display: flex;
    align-items: center;
    gap: 1rem;
}

.section-header::after {
    content: '';
    flex-grow: 1;
    height: 1px;
    background: linear-gradient(to right, rgba(255, 255, 255, 0.1), transparent);
}

/* Custom Metric Card */
.premium-metric {
    background: var(--glass-bg);
    border: 1px solid var(--glass-border);
    border-radius: 20px;
    padding: 1.5rem;
    text-align: center;
    transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
}

.premium-metric:hover {
    border-color: rgba(59, 130, 246, 0.4);
    transform: translateY(-5px);
    background: rgba(30, 41, 59, 0.6);
    box-shadow: 0 10px 25px -5px rgba(0, 0, 0, 0.3);
}

.metric-label {
    font-size: 0.75rem;
    text-transform: uppercase;
    letter-spacing: 0.05em;
    color: #94a3b8;
    margin-bottom: 0.5rem;
}

.metric-val {
    font-family: 'Outfit', sans-serif;
    font-size: 2.25rem;
    font-weight: 700;
    color: #fff;
    line-height: 1;
}

/* Policy Cards */
.policy-card {
    background: linear-gradient(90deg, rgba(30, 41, 59, 0.4) 0%, rgba(15, 23, 42, 0.4) 100%);
    border: 1px solid var(--glass-border);
    border-radius: 16px;
    padding: 1.25rem 1.75rem;
    margin-bottom: 1rem;
    transition: all 0.2s ease;
    display: flex;
    justify-content: space-between;
    align-items: center;
}

.policy-card:hover {
    border-color: rgba(59, 130, 246, 0.5);
    background: rgba(30, 41, 59, 0.6);
    transform: translateX(4px);
    box-shadow: 0 4px 20px -5px rgba(0, 0, 0, 0.3);
}

.card-title {
    font-family: 'Outfit', sans-serif;
    font-size: 1.15rem;
    font-weight: 700;
    color: #f8fafc;
    margin-bottom: 0.4rem;
}

.card-meta {
    font-size: 0.78rem;
    color: #94a3b8;
    display: flex;
    align-items: center;
    gap: 1.25rem;
}

.status-badge {
    padding: 3px 12px;
    border-radius: 20px;
    font-size: 0.65rem;
    font-weight: 800;
    letter-spacing: 0.05em;
    text-transform: uppercase;
}

.badge-binary { background: rgba(16, 185, 129, 0.1); color: #10b981; border: 1px solid rgba(16, 185, 129, 0.2); }
.badge-meta { background: rgba(245, 158, 11, 0.1); color: #f59e0b; border: 1px solid rgba(245, 158, 11, 0.2); }

/* Document Viewer Enhancement */
.viewer-container {
    background: linear-gradient(180deg, rgba(30, 41, 59, 0.4) 0%, rgba(15, 23, 42, 0.6) 100%);
    border: 1px solid var(--glass-border);
    border-radius: 24px;
    padding: 2rem;
    margin-top: 2rem;
    animation: slideUp 0.6s ease-out;
}

.meta-chip {
    padding: 6px 16px;
    background: rgba(255, 255, 255, 0.03);
    border: 1px solid var(--glass-border);
    border-radius: 30px;
    font-size: 0.75rem;
    font-weight: 600;
    color: #94a3b8;
    display: flex;
    align-items: center;
    gap: 8px;
}

.document-frame {
    border: 1px solid rgba(255, 255, 255, 0.1);
    border-radius: 16px;
    overflow: hidden;
    box-shadow: 0 25px 50px -12px rgba(0, 0, 0, 0.5);
    background: #000;
}

/* Tab/Reader Animation */
@keyframes slideUp {
    from { opacity: 0; transform: translateY(20px); }
    to { opacity: 1; transform: translateY(0); }
}

.tab-content { animation: slideUp 0.6s ease-out; }
</style>

<div class="page-hero">
    <div style="position: relative; z-index: 1; display: flex; align-items: flex-start; gap: 2rem;">
        <div style="font-size: 3.5rem; background: rgba(59, 130, 246, 0.1); width: 85px; height: 85px; display: flex; align-items: center; justify-content: center; border-radius: 20px; border: 1px solid rgba(59, 130, 246, 0.2); box-shadow: 0 0 20px rgba(59, 130, 246, 0.1);">📘</div>
        <div style="flex-grow: 1;">
            <h1 class="hero-title">Policy & Intelligence Guidance</h1>
            <p class="hero-subtitle">Unified frameworks for data handling, classification standards, and cross-functional accountability matrices.</p>
            <div style="display: flex; gap: 1rem; margin-top: 1.5rem;">
                <div style="background: rgba(16, 185, 129, 0.1); color: #10b981; padding: 5px 14px; border-radius: 30px; font-size: 0.78rem; font-weight: 600; display: flex; align-items: center; gap: 8px; border: 1px solid rgba(16, 185, 129, 0.2);">
                    <span style="width: 8px; height: 8px; border-radius: 50%; background: #10b981; box-shadow: 0 0 8px #10b981;"></span>
                    Operational Polices Verified
                </div>
                <div style="background: rgba(59, 130, 246, 0.1); color: #3b82f6; padding: 5px 14px; border-radius: 30px; font-size: 0.78rem; font-weight: 600; display: flex; align-items: center; gap: 8px; border: 1px solid rgba(59, 130, 246, 0.2);">
                    <span style="width: 8px; height: 8px; border-radius: 50%; background: #3b82f6; box-shadow: 0 0 8px #3b82f6;"></span>
                    V2.5 Governance Stack
                </div>
            </div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

render_quick_links()

# -----------------------------------------------------------------------------
# DATABASE RESOLUTION (Synced with Global Filters)
# -----------------------------------------------------------------------------
db_name = st.session_state.get('sf_database') or settings.SNOWFLAKE_DATABASE
if not db_name or str(db_name).upper() in ['NONE', '', 'NULL', 'UNKNOWN', '(NONE)']:
    db_name = "DATA_CLASSIFICATION_DB"

# Ensure Schema and Table exist
try:
    snowflake_connector.execute_non_query(f"CREATE SCHEMA IF NOT EXISTS {db_name}.DATA_CLASSIFICATION_GOVERNANCE")
    
    # Standardize Table Schema
    snowflake_connector.execute_non_query(f"""
        CREATE TABLE IF NOT EXISTS {db_name}.DATA_CLASSIFICATION_GOVERNANCE.POLICIES (
            POLICY_ID STRING NOT NULL DEFAULT UUID_STRING(),
            POLICY_NAME STRING NOT NULL,
            FILE_CONTENT BINARY NOT NULL,
            FILE_NAME STRING NOT NULL,
            MIME_TYPE STRING,
            FILE_SIZE NUMBER,
            POLICY_CONTENT TEXT, -- Extracted text for analysis
            CREATED_AT TIMESTAMP_NTZ DEFAULT CURRENT_TIMESTAMP(),
            CREATED_BY STRING DEFAULT CURRENT_USER(),
            UPDATED_AT TIMESTAMP_NTZ DEFAULT CURRENT_TIMESTAMP(),
            CONSTRAINT PK_POLICIES PRIMARY KEY (POLICY_ID)
        )
    """)
    # Ensure all target columns exist for older tables
    for col, ctype in [("FILE_CONTENT", "BINARY"), ("POLICY_CONTENT", "TEXT"), ("UPDATED_AT", "TIMESTAMP_NTZ")]:
        try:
            snowflake_connector.execute_non_query(f"ALTER TABLE {db_name}.DATA_CLASSIFICATION_GOVERNANCE.POLICIES ADD COLUMN IF NOT EXISTS {col} {ctype}")
        except Exception: pass
except Exception:
    pass

# One-time cleanup for requested removals
if 'cleanup_done' not in st.session_state:
    try:
        snowflake_connector.execute_non_query(
            f"DELETE FROM {db_name}.DATA_CLASSIFICATION_GOVERNANCE.POLICIES WHERE POLICY_NAME IN ('data classification', 'Data_classification_test')"
        )
        st.session_state['cleanup_done'] = True
    except Exception: pass

# -----------------------------------------------------------------------------
# DYNAMIC DATA FETCHING
# -----------------------------------------------------------------------------

def get_governance_roles():
    """Fetch Snowflake roles synchronized with application responsibility metadata."""
    try:
        sf_roles_raw = snowflake_connector.execute_query("SHOW ROLES") or []
        sf_role_names = [r['name'].upper() for r in sf_roles_raw]
        
        meta_rows = snowflake_connector.execute_query(
            f"SELECT ROLE_NAME, DESCRIPTION FROM {db_name}.DATA_GOVERNANCE.ROLES"
        ) or []
        meta_dict = {r['ROLE_NAME'].upper(): r['DESCRIPTION'] for r in meta_rows}
        
        sync_roles = []
        for name in sorted(sf_role_names):
            sync_roles.append({
                "ROLE": name,
                "RESPONSIBILITY": meta_dict.get(name, "General functional role; specific governance responsibility not defined.")
            })
        return sync_roles
    except Exception:
        return []

def get_role_counts():
    """Fetch counts per major governance category for the dashboard."""
    counts = {"Owners": 0, "Stewards": 0, "Analysts": 0, "Consumers": 0}
    try:
        query = f"SELECT ROLE_NAME, COUNT(DISTINCT USER_EMAIL) as CNT FROM {db_name}.DATA_GOVERNANCE.ROLE_ASSIGNMENTS GROUP BY 1"
        rows = snowflake_connector.execute_query(query) or []
        for r in rows:
            rn = str(r['ROLE_NAME']).upper()
            if "OWNER" in rn: counts["Owners"] += r['CNT']
            elif "STEWARD" in rn or "CUSTODIAN" in rn: counts["Stewards"] += r['CNT']
            elif "ANALYST" in rn or "SPECIALIST" in rn: counts["Analysts"] += r['CNT']
            else: counts["Consumers"] += r['CNT']
    except Exception:
        pass
    return counts

def get_key_policies():
    """Fetch all policies present in the POLICIES table."""
    policies = []
    try:
        # Check if table exists
        snowflake_connector.execute_query(f"SELECT 1 FROM {db_name}.DATA_CLASSIFICATION_GOVERNANCE.POLICIES LIMIT 1")
        
        query = f"""
            SELECT POLICY_ID, POLICY_NAME, POLICY_CONTENT, FILE_NAME, MIME_TYPE, 
                   FILE_CONTENT,
                   REGEXP_REPLACE(LEFT(POLICY_CONTENT, 200), '[\r\n\t]+', ' ', 1, 0) as SNIP
            FROM {db_name}.DATA_CLASSIFICATION_GOVERNANCE.POLICIES
            ORDER BY CREATED_AT DESC
        """
        rows = snowflake_connector.execute_query(query) or []
        for r in rows:
            raw_data = r.get('FILE_CONTENT')
            
            # CRITICAL FIX: Snowflake BINARY often returns as Hex strings
            if isinstance(raw_data, str) and len(raw_data) > 0:
                try:
                    # Attempt hex decode if it looks like hex (only hex chars)
                    # PDFs: 25504446, Word: 504B0304, but we'll be more generic
                    if all(c in '0123456789ABCDEFabcdef' for c in raw_data[:50].strip()):
                        raw_data = bytes.fromhex(raw_data.strip())
                except Exception:
                    pass
            
            # Ensure standard bytes format
            if isinstance(raw_data, (bytearray, memoryview)):
                raw_data = bytes(raw_data)
                
            policies.append({
                "TITLE": r.get('POLICY_NAME'),
                "VERSION": "Current", 
                "DESC": r.get('SNIP') or "No summary available",
                "FULL": r.get('POLICY_CONTENT'),
                "FNAME": r.get('FILE_NAME') or f"{r.get('POLICY_NAME')}.txt",
                "MIME": r.get('MIME_TYPE') or "text/plain",
                "RAW": raw_data
            })
    except Exception:
        pass
    
    return policies

# -----------------------------------------------------------------------------
# UI LAYOUT
# -----------------------------------------------------------------------------

tab_dash, tab_manage = st.tabs(["📊 Policy Dashboard", "📤 Upload & Manage"])

# --- DASHBOARD TAB ---
with tab_dash:
    # 2. Main Dashboard Content
    role_counts = get_role_counts()
    policies = get_key_policies()
    
    # Adjusted to a single column flow for better focus on policies
    policies = get_key_policies()
    
    if not policies:
        st.info("No documents found. Upload governance files to begin.")
    else:
        st.markdown('<div class="section-header">📦 Synchronized Governance Vault</div>', unsafe_allow_html=True)
        
        # Premium Card Loop for Dashboard
        for p in policies:
            status_class = "badge-binary" if p.get('RAW') else "badge-meta"
            status_label = "RAW BINARY" if p.get('RAW') else "META ONLY"
            
            st.markdown(f"""
<div class="policy-card">
    <div>
        <div class="card-title">{p['TITLE']}</div>
        <div class="card-meta">
            <span>📄 {p['FNAME']}</span>
            <span>🏷️ {p['MIME']}</span>
            <span class="status-badge {status_class}">{status_label}</span>
        </div>
        <div style="font-size: 0.85rem; color: #64748b; margin-top: 0.75rem; max-width: 600px;">
            {p['DESC']}...
        </div>
    </div>
    <div style="color: #3b82f6; font-size: 1.5rem; opacity: 0.5;">🔒</div>
</div>
""", unsafe_allow_html=True)
        
        # Interactive Policy Reader
        st.markdown('<div class="section-header">📖 INTELLECTUAL POLICY READER</div>', unsafe_allow_html=True)
        sel_pol_title = st.selectbox("Select Policy Document to Read or Download", options=[p['TITLE'] for p in policies])
        
        if sel_pol_title:
            vp = next(p for p in policies if p["TITLE"] == sel_pol_title)
            
            # TRIGGER LOGIC: Reset visibility if user switches policies in the dropdown
            active_vp_title = st.session_state.get('viewing_policy', {}).get('TITLE')
            if active_vp_title and active_vp_title != vp['TITLE']:
                st.session_state['viewing_policy_active'] = False

            p_c1, p_c2 = st.columns([1, 1])
            with p_c1:
                is_currently_viewing = st.session_state.get('viewing_policy_active') and st.session_state.get('viewing_policy', {}).get('TITLE') == vp['TITLE']
                
                if st.button("👁️ View Full Content" if not is_currently_viewing else "✅ Currently Viewing", key="btn_view_pol", width='stretch'):
                     st.session_state['viewing_policy'] = vp
                     st.session_state['viewing_policy_active'] = True
                     st.rerun()
            with p_c2:
                # Robust download selection: Raw Binary (Original) > Full Text Fallback
                raw_binary = vp.get('RAW')
                if isinstance(raw_binary, (bytearray, memoryview)):
                    raw_binary = bytes(raw_binary)
                    
                dl_data = raw_binary if raw_binary else (vp.get('FULL') or "")
                
                st.download_button(
                    label="📥 Download Policy" if not raw_binary else "📥 Download Original PDF/Doc",
                    data=dl_data,
                    file_name=vp.get('FNAME') or f"{vp.get('TITLE').replace(' ', '_')}.txt",
                    mime=vp.get('MIME') or "text/plain",
                    key=f"dl_dash_{vp.get('TITLE')}",
                    width='stretch',
                    help="Downloads the original uploaded binary file if available, otherwise falls back to extracted text."
                )

        # Show Content Only if active for the currently selected document
        if st.session_state.get('viewing_policy_active') and st.session_state.get('viewing_policy', {}).get('TITLE') == sel_pol_title:
            vp = st.session_state['viewing_policy']
            
            st.markdown(f"""
<div class="viewer-container">
    <div class="section-header">📜 AUTHORITATIVE SOURCE VIEW</div>
    <h2 style="font-family: 'Outfit', sans-serif; color: #fff; margin: 0 0 1.5rem 0; font-size: 1.75rem;">{vp["TITLE"]}</h2>
    <div style="display: flex; gap: 1rem; margin-bottom: 2.5rem;">
        <div class="meta-chip">📁 {vp.get('FNAME')}</div>
        <div class="meta-chip">🏷️ {vp.get('MIME')}</div>
        <div class="meta-chip">
            <span style="width: 8px; height: 8px; border-radius: 50%; background: #3b82f6; box-shadow: 0 0 8px #3b82f6;"></span>
            Verified Stream
        </div>
    </div>
""", unsafe_allow_html=True)
            
            # 1. Original PDF Rendering (Full Image)
            if vp.get('RAW') and (vp.get('MIME') == 'application/pdf' or str(vp.get('FNAME', '')).lower().endswith('.pdf')):
                try:
                    raw_data = vp['RAW']
                    if isinstance(raw_data, str) and all(c in '0123456789ABCDEFabcdef' for c in raw_data[:20]):
                         raw_data = bytes.fromhex(raw_data.strip())
                    
                    if isinstance(raw_data, bytes):
                        b64_pdf = base64.b64encode(raw_data).decode('utf-8')
                        pdf_display = f"""
                        <div class="document-frame">
                        <object data="data:application/pdf;base64,{b64_pdf}" type="application/pdf" width="100%" height="1200">
                            <iframe src="data:application/pdf;base64,{b64_pdf}" width="100%" height="1200" style="border:none;">
                                <p>PDF Display fallback: <a href="data:application/pdf;base64,{b64_pdf}" download="{vp.get('FNAME', 'policy.pdf')}">Download the PDF</a></p>
                            </iframe>
                        </object>
                        </div>
                        """
                        st.markdown(pdf_display, unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"Render Error: {e}")
            else:
                st.info("📂 Inline preview is currently optimized for PDF documents. Please use the Download button for other file types.")
            
            st.markdown("<br><br>", unsafe_allow_html=True)
            
            # 3. Action Footer
            c1, c2, c3 = st.columns([1, 1, 2])
            with c1:
                if st.button("❌ Close Review", key="close_render_main", width='stretch'):
                    if 'viewing_policy' in st.session_state: del st.session_state['viewing_policy']
                    if 'viewing_policy_active' in st.session_state: del st.session_state['viewing_policy_active']
                    st.rerun()
            with c2:
                if st.button("🔄 Force Refresh", key="reload_render", width='stretch'):
                    st.rerun()
            
            st.markdown("</div>", unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown('<div class="section-header">📐 TRUST TIERS & CLASSIFICATION BANDS</div>', unsafe_allow_html=True)
    
    with st.container(border=True):
        f1, f2, f3 = st.columns(3)
        
        with f1:
            st.markdown("**Confidentiality (C)**")
            tiers = [("C0 - Public", "#10b981"), ("C1 - Internal", "#3b82f6"), ("C2 - Restricted", "#f59e0b"), ("C3 - Confidential", "#f43f5e")]
            for t, c in tiers:
                st.markdown(f'<div style="display:flex; align-items:center; gap:10px; margin-bottom:8px;"><div style="width:10px; height:10px; border-radius:50%; background:{c}; box-shadow: 0 0 8px {c}40;"></div><span style="font-size:0.85rem; color:#cbd5e1;">{t}</span></div>', unsafe_allow_html=True)
        
        with f2:
            st.markdown("**Integrity (I)**")
            tiers = [("I0 - Low", "#64748b"), ("I1 - Standard", "#3b82f6"), ("I2 - High", "#8b5cf6"), ("I3 - Critical", "#f43f5e")]
            for t, c in tiers:
                st.markdown(f'<div style="display:flex; align-items:center; gap:10px; margin-bottom:8px;"><div style="width:10px; height:10px; border-radius:50%; background:{c}; box-shadow: 0 0 8px {c}40;"></div><span style="font-size:0.85rem; color:#cbd5e1;">{t}</span></div>', unsafe_allow_html=True)
        
        with f3:
            st.markdown("**Availability (A)**")
            tiers = [("A0 - Low", "#64748b"), ("A1 - Standard", "#3b82f6"), ("A2 - High", "#8b5cf6"), ("A3 - Critical", "#f43f5e")]
            for t, c in tiers:
                st.markdown(f'<div style="display:flex; align-items:center; gap:10px; margin-bottom:8px;"><div style="width:10px; height:10px; border-radius:50%; background:{c}; box-shadow: 0 0 8px {c}40;"></div><span style="font-size:0.85rem; color:#cbd5e1;">{t}</span></div>', unsafe_allow_html=True)

# --- UPLOAD & MANAGE TAB ---
with tab_manage:
    st.markdown(f"""
<div style="background: linear-gradient(90deg, rgba(16, 185, 129, 0.08), rgba(0, 0, 0, 0)); padding: 15px; border-radius: 12px; border-left: 3px solid #10b981; margin-bottom: 25px;">
    <h3 style="margin:0; color:white; font-size:1.3rem;">📤 Policy Lifecycle Management</h3>
    <p style="margin:5px 0 0 0; color:rgba(255,255,255,0.6); font-size:0.9rem;">
        Maintain authoritative policy versions and business rules. Use the <b>AI Document Parser</b> for intelligent field extraction.
    </p>
</div>
""", unsafe_allow_html=True)
    
    # Schema and Table managed in initialization block above
    pass

    action = st.radio("Management Phase", ["🆕 Create New Policy", "🔄 Update Existing Documentation"], horizontal=True)

    # 2. Fetch existing for dropdown
    existing_policies = []
    try:
        existing_policies = snowflake_connector.execute_query(
            f"SELECT * FROM {db_name}.DATA_CLASSIFICATION_GOVERNANCE.POLICIES WHERE POLICY_NAME IS NOT NULL ORDER BY POLICY_NAME"
        ) or []
    except Exception: pass

    default_title = ""
    default_content = ""
    
    if "Update" in action:
        if not existing_policies:
            st.warning("No existing policies found.")
        else:
            pol_map = {f"{p['POLICY_NAME']}": p for p in existing_policies}
            sel_pol_label = st.selectbox("Select Policy to Update", list(pol_map.keys()))
            if sel_pol_label:
                sel_pol = pol_map[sel_pol_label]
                selected_policy_id = sel_pol['POLICY_ID']
                default_title = sel_pol['POLICY_NAME']
                default_content = sel_pol.get('POLICY_CONTENT', "")
                
                try:
                    v_parts = str(sel_pol.get('POLICY_VERSION', "1.0")).split('.')
                    if len(v_parts) == 2: default_ver = f"{v_parts[0]}.{int(v_parts[1])+1}"
                    else: default_ver = "1.1"
                except: default_ver = "1.1"

    with st.container(border=True):
        st.markdown("#### 📝 Policy Details")
        with st.form("policy_form", border=False):
            f_col1, f_col2 = st.columns(2)
            with f_col1:
                p_name = st.text_input("Policy Title / Name", value=default_title, disabled=("Update" in action), help="Authoritative name of the policy.")
                uploaded_file = st.file_uploader("Upload Document (PDF/DOCX/TXT)", type=['pdf','docx','doc','txt','md'])
            with f_col2:
                p_content = st.text_area("Analysis / Strategic Notes", value=default_content, height=180, help="Keywords, scope, or manual summary.")
            
            st.markdown("<br>", unsafe_allow_html=True)
            submit = st.form_submit_button("🚀 Commit Policy to Governance Vault", width='stretch', type="primary")
    
    if submit:
        final_content = p_content
        fname = None
        fsize = 0
        fmime = "text/plain"
        
        if uploaded_file:
            fname = uploaded_file.name
            fsize = uploaded_file.size
            fmime = uploaded_file.type
            try:
                raw_bytes = uploaded_file.read()
                # More robust PDF detection
                is_pdf = fmime == "application/pdf" or fname.lower().endswith(".pdf") or raw_bytes.startswith(b"%PDF")
                
                if is_pdf:
                    extracted_text = ""
                    # Try pypdf first (modern), then PyPDF2 (requested)
                    try:
                        if pypdf:
                            reader = pypdf.PdfReader(io.BytesIO(raw_bytes))
                            extracted_text = "\n\n".join([p.extract_text() or "" for p in reader.pages]).strip()
                        elif PyPDF2:
                            # Handling older version vs newer v3 API compatibility
                            reader = PyPDF2.PdfReader(io.BytesIO(raw_bytes))
                            extracted_text = "\n\n".join([p.extract_text() or "" for p in reader.pages]).strip()
                            
                        if extracted_text:
                            final_content = extracted_text
                            st.info(f"✅ Extracted text from PDF using {'pypdf' if pypdf else 'PyPDF2'}")
                        else:
                            final_content = "[System: PDF extraction yielded no text. The document may be scanned or image-based.]"
                    except Exception as pdf_err:
                        final_content = f"[System: PDF Processing Error: {pdf_err}]"
                elif fname.lower().endswith((".docx", ".doc")):
                    if docx:
                        try:
                            doc = docx.Document(io.BytesIO(raw_bytes))
                            final_content = "\n".join([para.text for para in doc.paragraphs]).strip()
                            st.info("✅ Extracted text from Word document.")
                        except Exception as doc_err:
                            final_content = f"[System: Word Processing Error: {doc_err}]"
                    else:
                        final_content = "[System: python-docx not available for extraction. Storing as raw binary metadata.]"
                else:
                    try:
                        final_content = raw_bytes.decode("utf-8", errors="ignore")
                    except Exception:
                        final_content = "[System: Binary or incompatible text encoding detected.]"
            except Exception as e:
                st.error(f"Error reading file: {e}")
                st.stop()
        
        if not p_name:
            st.error("Policy Name is required.")
        elif not final_content and not p_content:
            st.error("Policy content is required (either via upload or manual entry).")
        else:
            save_content = final_content if uploaded_file else p_content
            import uuid
            try:
                # Ensure we are passing standard bytes to the connector for storage
                input_bytes = raw_bytes if uploaded_file else None
                if isinstance(input_bytes, bytearray):
                    input_bytes = bytes(input_bytes)

                if "🆕" in action:
                    snowflake_connector.execute_non_query(
                        f"INSERT INTO {db_name}.DATA_CLASSIFICATION_GOVERNANCE.POLICIES (POLICY_NAME, FILE_CONTENT, FILE_NAME, MIME_TYPE, FILE_SIZE, POLICY_CONTENT, CREATED_AT, CREATED_BY) VALUES (%(name)s, %(file)s, %(fn)s, %(fmt)s, %(fs)s, %(cont)s, CURRENT_TIMESTAMP, CURRENT_USER)",
                        {"name": p_name, "file": input_bytes, "fn": fname, "fmt": fmime, "fs": fsize, "cont": save_content}
                    )
                    st.success(f"🚀 Successfully stored policy: {p_name}. Raw binary preserved for retrieval.")
                else: 
                    # SELECTIVE UPDATE: Only update FILE_CONTENT if a new file was actually uploaded
                    if selected_policy_id:
                        if uploaded_file:
                            snowflake_connector.execute_non_query(
                                f"UPDATE {db_name}.DATA_CLASSIFICATION_GOVERNANCE.POLICIES SET POLICY_NAME=%(name)s, FILE_CONTENT=%(file)s, FILE_NAME=%(fn)s, MIME_TYPE=%(fmt)s, FILE_SIZE=%(fs)s, POLICY_CONTENT=%(cont)s, UPDATED_AT=CURRENT_TIMESTAMP WHERE POLICY_ID=%(id)s",
                                {"id": selected_policy_id, "name": p_name, "file": input_bytes, "fn": fname, "fmt": fmime, "fs": fsize, "cont": save_content}
                            )
                            st.success(f"🔄 Updated record and file: {p_name}")
                        else:
                            snowflake_connector.execute_non_query(
                                f"UPDATE {db_name}.DATA_CLASSIFICATION_GOVERNANCE.POLICIES SET POLICY_NAME=%(name)s, POLICY_CONTENT=%(cont)s, UPDATED_AT=CURRENT_TIMESTAMP WHERE POLICY_ID=%(id)s",
                                {"id": selected_policy_id, "name": p_name, "cont": save_content}
                            )
                            st.success(f"📝 Updated metadata/notes for: {p_name} (Existing binary file kept)")
                    else: st.error("Policy ID missing for update.")
                
                time.sleep(1)
                st.rerun()
            except Exception as e: st.error(f"Database Error: {e}")

    # --- BULK MAINTENANCE SECTION ---
    st.markdown("<br>", unsafe_allow_html=True)
    with st.expander("🛠️ Policy Database Maintenance"):
        st.write("Extract text content from existing binary records in the table using **PyPDF2** and **python-docx**.")
        if st.button("🔍 Scan and Extract Missing Content"):
            try:
                # Find policies with raw data but generic/missing content
                # Search FILE_CONTENT which is the new binary home
                missing = snowflake_connector.execute_query(
                    f"SELECT POLICY_ID, POLICY_NAME, MIME_TYPE, FILE_NAME, FILE_CONTENT FROM {db_name}.DATA_CLASSIFICATION_GOVERNANCE.POLICIES WHERE FILE_CONTENT IS NOT NULL AND (POLICY_CONTENT IS NULL OR POLICY_CONTENT LIKE '[System:%')"
                ) or []
                
                if not missing:
                    st.success("No records require maintenance.")
                else:
                    progress_bar = st.progress(0)
                    for i, row in enumerate(missing):
                        row_id = row['POLICY_ID']
                        raw_data = row['FILE_CONTENT']
                        
                        # Handle Snowflake Hex String return format for extraction
                        if isinstance(raw_data, str) and raw_data.startswith(("25504446", "504B0304")):
                            try:
                                raw_data = bytes.fromhex(raw_data)
                            except Exception: pass
                        elif isinstance(raw_data, bytearray):
                            raw_data = bytes(raw_data)

                        mime = row['MIME_TYPE'] or ""
                        fname = row['FILE_NAME'] or ""
                        new_content = None
                        
                        try:
                            # PDF extraction
                            if 'pdf' in mime.lower() or fname.lower().endswith('.pdf'):
                                if pypdf or PyPDF2:
                                    reader = pypdf.PdfReader(io.BytesIO(raw_data)) if pypdf else PyPDF2.PdfReader(io.BytesIO(raw_data))
                                    new_content = "\n\n".join([p.extract_text() or "" for p in reader.pages]).strip()
                            
                            # Word extraction
                            elif 'word' in mime.lower() or 'officedocument' in mime.lower() or fname.lower().endswith(('.docx', '.doc')):
                                if docx:
                                    doc = docx.Document(io.BytesIO(raw_data))
                                    new_content = "\n".join([para.text for para in doc.paragraphs]).strip()
                            
                            if new_content:
                                snowflake_connector.execute_non_query(
                                    f"UPDATE {db_name}.DATA_CLASSIFICATION_GOVERNANCE.POLICIES SET POLICY_CONTENT=%(cont)s, UPDATED_AT=CURRENT_TIMESTAMP WHERE POLICY_ID=%(id)s",
                                    {"cont": new_content, "id": row_id}
                                )
                                st.write(f"✅ Extracted: **{row['POLICY_NAME']}**")
                            else:
                                st.write(f"⚠️ Could not extract from: **{row['POLICY_NAME']}**")
                                
                        except Exception as row_err:
                            st.error(f"Error on {row['POLICY_NAME']}: {row_err}")
                            
                        progress_bar.progress((i + 1) / len(missing))
                    
                    st.success(f"Maintenance complete. Processed {len(missing)} records.")
                    st.rerun()
            except Exception as e:
                st.error(f"Maintenance Scan Error: {e}")
    # List all uploaded documents in the Manage tab for verification
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown('<div class="section-header">📦 Current Governance Inventory</div>', unsafe_allow_html=True)
    
    inventory_policies = get_key_policies()
    if not inventory_policies:
        st.info("No documents uploaded yet.")
    else:
        # Mini Card View for Inventory
        for p in inventory_policies:
            status_class = "badge-binary" if p.get('RAW') else "badge-meta"
            status_label = "STABLE" if p.get('RAW') else "TEXT ONLY"
            
            st.markdown(f"""
<div class="policy-card" style="padding: 1rem 1.5rem; border-radius: 12px; margin-bottom: 0.75rem;">
    <div>
        <div style="font-weight: 700; color: #fff; font-size: 0.95rem;">{p['TITLE']}</div>
        <div style="font-size: 0.75rem; color: #64748b; margin-top: 2px;">
            {p['FNAME']} • {p['MIME']}
        </div>
    </div>
    <div class="status-badge {status_class}">{status_label}</div>
</div>
""", unsafe_allow_html=True)
